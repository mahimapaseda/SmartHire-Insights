import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Ensure output cv directory exists
CV_DIR = "cv"
os.makedirs(CV_DIR, exist_ok=True)

# Define candidates data with standard Black & White styling
CANDIDATES = [
    {
        "name": "Nimesh Perera",
        "email": "nimesh.perera92@gmail.com",
        "phone": "+94 71 245 6731",
        "role": "Software Engineer",
        "company": "WSO2",
        "duration": "2 years (2024 - Present)",
        "degree": "BSc (Hons) in Software Engineering",
        "inst": "NSBM Green University",
        "skills": ["Java", "Spring Boot", "React", "MySQL", "Git"],
        "summary": "Results-oriented Software Engineer with a solid foundation in developing enterprise-grade cloud-native applications. Proficient in Java, Spring Boot microservices, and React frontend structures. Passionate about writing clean, testable, and highly efficient code to drive system scalability.",
        "experience_bullets": [
            "Designed and implemented high-throughput REST APIs using Spring Boot and Java, improving server response times by 25%.",
            "Developed rich, interactive UI components using React, Redux, and modern styling utilities, enhancing overall user engagement.",
            "Collaborated with cross-functional Agile teams to deliver key system updates on time, participating in continuous code reviews.",
            "Optimized SQL databases and structured query patterns in MySQL to improve transaction speeds and ensure seamless indexing.",
            "Leveraged Git for robust version control, establishing seamless branch management and CI/CD pipelines via GitHub actions."
        ],
        "edu_details": "BSc (Hons) in Software Engineering – Graduated with First Class Honours. Active member of the University Computing Society and competitive programming team."
    },
    {
        "name": "Kavindi Senanayake",
        "email": "kavindisena@gmail.com",
        "phone": "+94 77 561 2284",
        "role": "Associate Software Engineer",
        "company": "99X",
        "duration": "1.5 years (2024 - Present)",
        "degree": "BSc in Information Technology",
        "inst": "University of Moratuwa",
        "skills": ["Python", "Django", "REST APIs", "PostgreSQL", "JavaScript"],
        "summary": "Enthusiastic and adaptable software developer specializing in Python and Django backend infrastructure. Proven track record of designing reliable, RESTful APIs and managing complex database architectures at 99X. Experienced in Moratuwa's rigorous technical environment.",
        "experience_bullets": [
            "Developed scalable backend workflows and REST API endpoints utilizing Django and Python, ensuring high reliability.",
            "Collaborated in architectural discussions to refactor legacy Python modules, achieving a 15% reduction in CPU utilization.",
            "Designed schema migrations and query optimizations in PostgreSQL, managing high volumes of relational candidate records.",
            "Participated actively in automated integration testing and unit testing to secure high quality code coverage.",
            "Contributed to dynamic frontend modules using Vanilla JavaScript and modern UI libraries."
        ],
        "edu_details": "BSc in Information Technology – Faculty of Information Technology, University of Moratuwa. Specialized in Database Systems and Distributed Applications."
    },
    {
        "name": "Tharindu Jayasuriya",
        "email": "tharinduj.dev@gmail.com",
        "phone": "+94 76 884 1902",
        "role": "Full Stack Developer",
        "company": "Sysco LABS",
        "duration": "3 years (2023 - Present)",
        "degree": "BSc in Computer Science",
        "inst": "SLIIT",
        "skills": ["Node.js", "Angular", "MongoDB", "Docker", "Express", "AWS"],
        "summary": "Versatile Full Stack Developer with 3 years of hands-on experience architecting high-performance enterprise systems at Sysco LABS. Expert in Node.js server frameworks, Angular single-page applications, and containerized Docker deployments.",
        "experience_bullets": [
            "Architected dynamic and highly reusable UI interfaces in Angular, cutting overall rendering latency by 30%.",
            "Constructed modular microservices using Node.js and Express, supporting robust real-time candidate search capabilities.",
            "Managed NoSQL database patterns using MongoDB, building robust indexing schemas for swift horizontal scale.",
            "Containerized and deployed backend and frontend applications using Docker, streamlining developer onboarding pipelines.",
            "Maintained deployment instances and automated server configurations using AWS EC2 and S3 buckets."
        ],
        "edu_details": "BSc in Computer Science – Sri Lanka Institute of Information Technology. Major in Software Engineering. Received academic excellence awards in final year project."
    },
    {
        "name": "Dilmi Fernando",
        "email": "dilmi.qa@gmail.com",
        "phone": "+94 71 993 4210",
        "role": "QA Engineer",
        "company": "Virtusa",
        "duration": "2 years (2024 - Present)",
        "degree": "BSc in IT",
        "inst": "Informatics Institute of Technology (IIT)",
        "skills": ["Selenium", "TestNG", "Manual Testing", "JIRA", "Java", "SQL"],
        "summary": "Detail-oriented QA Engineer with 2 years of extensive automation and manual testing experience at Virtusa. Highly skilled in designing robust Selenium test automation frameworks and managing detailed bug lifecycles in fast-paced Agile environments.",
        "experience_bullets": [
            "Created and optimized robust, scalable automation test suites using Selenium WebDriver and TestNG, increasing test coverage by 40%.",
            "Led detailed manual testing efforts, creating high-level test plans, comprehensive test scenarios, and regression reports.",
            "Managed defect life cycles, sprint tracking, and quality assurance metrics within JIRA, aligning directly with business needs.",
            "Conducted API testing using Postman to verify payload structures, HTTP status codes, and JSON response bodies.",
            "Formulated complex SQL queries to validate backend database states and ensure data integrity during schema updates."
        ],
        "edu_details": "BSc (Hons) in Information Technology – Informatics Institute of Technology (IIT), affiliated with the University of Westminster, UK. Specialized in Software Quality Assurance."
    },
    {
        "name": "Rukshan Maduranga",
        "email": "rukshan.maduranga@gmail.com",
        "phone": "+94 75 118 5642",
        "role": "Software QA Analyst",
        "company": "CodeGen",
        "duration": "1 year (2025 - Present)",
        "degree": "BSc in Software Engineering",
        "inst": "General Sir John Kotelawala Defence University (KDU)",
        "skills": ["Automation Testing", "Cypress", "API Testing", "Postman", "JavaScript"],
        "summary": "Highly motivated Software QA Analyst with a strong background in Software Engineering from KDU. Proven technical acumen in modern end-to-end automation testing frameworks like Cypress, with a strong focus on high-fidelity API testing and validation.",
        "experience_bullets": [
            "Authored high-performance end-to-end web integration test scripts in JavaScript using Cypress, reducing release times.",
            "Conducted regression and performance checks on new microservices interfaces prior to production rollouts.",
            "Developed fully automated API verification tests in Postman, ensuring schema correctness for complex request bodies.",
            "Partnered with development teams in CodeGen to identify edge-case defects, ensuring rapid resolution before release.",
            "Participated actively in daily standups, review boards, and sprint planning sessions in an Agile ecosystem."
        ],
        "edu_details": "BSc (Hons) in Software Engineering – Faculty of Computing, General Sir John Kotelawala Defence University (KDU). Focused on Software Architectures and Agile Methodologies."
    },
    {
        "name": "Shenali Wickramasinghe",
        "email": "shenaliwick.qa@gmail.com",
        "phone": "+94 77 662 9081",
        "role": "QA Associate",
        "company": "IFS",
        "duration": "2.5 years (2023 - Present)",
        "degree": "BSc in Computing",
        "inst": "University of Colombo School of Computing (UCSC)",
        "skills": ["Regression Testing", "Load Testing", "SQL", "JMeter", "GitHub"],
        "summary": "Meticulous QA Associate with 2.5 years of experience at IFS. UCSC graduate with deep expertise in enterprise regression testing, database validation using SQL, and performance/load testing using Apache JMeter. Committed to maximizing software excellence.",
        "experience_bullets": [
            "Conducted rigorous regression testing on core enterprise resource planning (ERP) modules, preventing critical regressions.",
            "Designed and executed performance and stress test plans using Apache JMeter, finding throughput limits.",
            "Wrote comprehensive SQL queries to trace database constraints, verify column states, and perform data audits.",
            "Tracked test scripts and shared source repositories using GitHub, establishing clean workflow integrations.",
            "Created high-fidelity documentation, including test specs, defect analysis reports, and user-acceptance test scripts."
        ],
        "edu_details": "BSc in Computing – University of Colombo School of Computing (UCSC). Active participant in university computing workshops and technology hackathons."
    },
    {
        "name": "Chamod Silva",
        "email": "chamodsilva.net@gmail.com",
        "phone": "+94 71 458 7743",
        "role": "Network Engineer",
        "company": "Dialog Axiata",
        "duration": "4 years (2022 - Present)",
        "degree": "BEng in Network Engineering",
        "inst": "University of Kelaniya",
        "skills": ["Cisco Routing", "Switching", "Firewall Management", "VPN", "Wireshark"],
        "summary": "Highly certified Network Engineer with 4 years of hands-on experience designing and managing core telecom infrastructure at Dialog Axiata. Expert in enterprise Cisco routing, switching, security configurations, and advanced firewall administration.",
        "experience_bullets": [
            "Managed, configured, and optimized enterprise Cisco routers and switches for high-availability core telecom operations.",
            "Implemented and audited advanced firewall policies and VPN setups, securing internal corporate communications.",
            "Analyzed network anomalies, packet traces, and routing loops using Wireshark, maintaining 99.99% system uptime.",
            "Led a major network virtualization migration project, improving internal server bandwidth distribution by 35%.",
            "Documented complex physical and logical network topologies, security protocols, and operational recovery plans."
        ],
        "edu_details": "BEng (Hons) in Network Engineering – University of Kelaniya. Focus on Enterprise Networks, Network Protocols, and Wireless Communication Architectures."
    },
    {
        "name": "Isuru Hettiarachchi",
        "email": "isuru.hetti@gmail.com",
        "phone": "+94 76 225 1988",
        "role": "IT Network Executive",
        "company": "Sri Lanka Telecom",
        "duration": "2 years (2024 - Present)",
        "degree": "Higher Diploma in Networking",
        "inst": "ESOFT Metro Campus",
        "skills": ["CCNA", "LAN/WAN", "Network Security", "DNS/DHCP", "Linux"],
        "summary": "Proactive IT Network Executive with 2 years of experience at Sri Lanka Telecom. CCNA certified professional with deep expertise configuring local area and wide area networks (LAN/WAN), resolving network security alerts, and administering Linux server structures.",
        "experience_bullets": [
            "Configured, tested, and resolved hardware issues on router systems, switches, and edge security devices.",
            "Managed corporate LAN/WAN configurations, optimizing IP subnets, VLAN allocations, and address pools.",
            "Monitored network traffic and enforced network security measures, resolving unauthorized access attempts.",
            "Maintained critical core infrastructure services, including DNS configurations, DHCP scopes, and dynamic routing.",
            "Administered Linux-based domain servers and internal corporate services, writing automation scripts in Bash."
        ],
        "edu_details": "Higher Diploma in Networking – ESOFT Metro Campus. Completed CCNA modules with distinction. Certified Cisco Network Associate."
    },
    {
        "name": "Piumi Ranasinghe",
        "email": "piumir.design@gmail.com",
        "phone": "+94 77 902 1135",
        "role": "Junior Fashion Designer",
        "company": "Hameedia",
        "duration": "2 years (2024 - Present)",
        "degree": "Diploma in Fashion Design",
        "inst": "Academy of Design (AOD)",
        "skills": ["Adobe Illustrator", "Textile Design", "Styling", "Pattern Making", "Trend Analysis"],
        "summary": "Creative and visionary Junior Fashion Designer with 2 years of hands-on design and styling experience at Hameedia. Proficient in digital illustration using Adobe Illustrator, pattern construction, textile color theory, and high-fidelity trend research.",
        "experience_bullets": [
            "Created trend-setting menswear and formalwear design collections, bringing concepts from mood boards to production.",
            "Developed intricate digital design patterns and technical tech packs using Adobe Illustrator and Photoshop.",
            "Led visual styling for local high-end product shoots, raising catalog conversion rates by 20%.",
            "Collaborated directly with structural pattern makers to verify proper garment fits and fabric choices.",
            "Conducted thorough global trend analysis to conceptualize seasonal mood colors and design palettes."
        ],
        "edu_details": "Diploma in Fashion Design – Academy of Design (AOD). Showcased graduation collection at the Sri Lanka Design Festival (SLDF). Specialized in Technical Apparel Design."
    },
    {
        "name": "Maleesha Karunaratne",
        "email": "maleeshafashion@gmail.com",
        "phone": "+94 71 770 4421",
        "role": "Fashion Coordinator",
        "company": "MAS Holdings",
        "duration": "3 years (2023 - Present)",
        "degree": "BA in Fashion Design",
        "inst": "Raffles Design Institute",
        "skills": ["Garment Construction", "Fashion Sketching", "Trend Analysis", "Product Lifecycle Management", "Tech Packs"],
        "summary": "Dynamic and meticulous Fashion Coordinator with 3 years of apparel development experience at MAS Holdings. Expert in managing the product lifecycle, orchestrating raw material sourcing, compiling technical specs, and interpreting international fashion movements.",
        "experience_bullets": [
            "Orchestrated cross-functional apparel product development cycles for leading global activewear accounts at MAS.",
            "Authored and audited detailed tech pack documentation, ensuring rigorous alignment with international specifications.",
            "Conducted seasonal fashion sketching and research on premium smart fabrics to support technical apparel lines.",
            "Coordinated bulk production runs, managing material sourcing lead times and ensuring on-time shipments.",
            "Evaluated structural garment construction samples, providing constructive feedback on fit, drape, and stitch metrics."
        ],
        "edu_details": "BA (Hons) in Fashion Design – Raffles Design Institute. Awarded 'Most Innovative Collection' in final runway show. Focused on Activewear Technology and Product Lifecycle Management."
    }
]

def add_horizontal_line(paragraph, color_hex="000000", size=12):
    """Inserts a solid black horizontal line below a paragraph via XML manipulation."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_docx(candidate):
    name = candidate["name"]
    filename = os.path.join(CV_DIR, f"{name.lower().replace(' ', '_')}_cv.docx")
    
    # Initialize Word Document
    doc = Document()
    
    # Page setup (Standard 1.0 inch margins)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Configure default font to Calibri (Standard professional corporate font)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0) # Strictly Black
    
    # Title/Header Table (2 columns: Left for Name/Title, Right for Contact)
    # Using a borderless table for clean alignments
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    
    # Set cell widths
    header_table.columns[0].width = Inches(4.5)
    header_table.columns[1].width = Inches(2.0)
    
    cell_left = header_table.cell(0, 0)
    cell_right = header_table.cell(0, 1)
    
    # Left Cell: Name and Target Role
    p_name = cell_left.paragraphs[0]
    run_name = p_name.add_run(name)
    run_name.font.size = Pt(24)
    run_name.bold = True
    run_name.font.color.rgb = RGBColor(0, 0, 0) # Strictly Black
    p_name.paragraph_format.space_after = Pt(2)
    
    p_role = cell_left.add_paragraph()
    run_role = p_role.add_run(candidate["role"])
    run_role.font.size = Pt(13)
    run_role.italic = True
    run_role.font.color.rgb = RGBColor(80, 80, 80) # Dark Neutral Gray
    p_role.paragraph_format.space_after = Pt(0)
    
    # Right Cell: Contact details (Right aligned)
    p_contact = cell_right.paragraphs[0]
    p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_email = p_contact.add_run(f"{candidate['email']}\n{candidate['phone']}")
    run_email.font.size = Pt(9.5)
    run_email.font.color.rgb = RGBColor(80, 80, 80) # Dark Neutral Gray
    p_contact.paragraph_format.space_after = Pt(0)
    
    # Bottom black line below header table
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(12)
    add_horizontal_line(p_line, color_hex="000000", size=12) # Pure black divider
    
    # SECTION 1: Professional Summary
    h_summary = doc.add_paragraph()
    run_h_summary = h_summary.add_run("PROFESSIONAL SUMMARY")
    run_h_summary.font.size = Pt(13)
    run_h_summary.bold = True
    run_h_summary.font.color.rgb = RGBColor(0, 0, 0) # Strictly Black
    h_summary.paragraph_format.space_before = Pt(10)
    h_summary.paragraph_format.space_after = Pt(4)
    
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(14)
    p_summary.paragraph_format.line_spacing = 1.15
    run_p_summary = p_summary.add_run(candidate["summary"])
    run_p_summary.font.color.rgb = RGBColor(0, 0, 0)
    
    # SECTION 2: Technical Skills
    h_skills = doc.add_paragraph()
    run_h_skills = h_skills.add_run("CORE STRENGTHS & SKILLS")
    run_h_skills.font.size = Pt(13)
    run_h_skills.bold = True
    run_h_skills.font.color.rgb = RGBColor(0, 0, 0) # Strictly Black
    h_skills.paragraph_format.space_before = Pt(10)
    h_skills.paragraph_format.space_after = Pt(6)
    
    p_skills = doc.add_paragraph()
    p_skills.paragraph_format.space_after = Pt(14)
    p_skills.paragraph_format.line_spacing = 1.2
    
    # Add skills with stylish bullet points separating tags in black
    skills_text = "  ·  ".join(candidate["skills"])
    run_skills = p_skills.add_run(skills_text)
    run_skills.bold = True
    run_skills.font.color.rgb = RGBColor(0, 0, 0)
    run_skills.font.size = Pt(11)
    
    # SECTION 3: Professional Experience
    h_exp = doc.add_paragraph()
    run_h_exp = h_exp.add_run("PROFESSIONAL EXPERIENCE")
    run_h_exp.font.size = Pt(13)
    run_h_exp.bold = True
    run_h_exp.font.color.rgb = RGBColor(0, 0, 0) # Strictly Black
    h_exp.paragraph_format.space_before = Pt(10)
    h_exp.paragraph_format.space_after = Pt(6)
    
    # Job Header Table (Title on Left, Company & Duration on Right)
    job_table = doc.add_table(rows=1, cols=2)
    job_table.autofit = False
    job_table.columns[0].width = Inches(4.5)
    job_table.columns[1].width = Inches(2.0)
    
    cell_j_left = job_table.cell(0, 0)
    cell_j_right = job_table.cell(0, 1)
    
    p_j_left = cell_j_left.paragraphs[0]
    run_j_title = p_j_left.add_run(candidate["role"])
    run_j_title.bold = True
    run_j_title.font.size = Pt(12)
    run_j_title.font.color.rgb = RGBColor(0, 0, 0)
    p_j_left.paragraph_format.space_after = Pt(0)
    
    p_j_right = cell_j_right.paragraphs[0]
    p_j_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_j_comp = p_j_right.add_run(f"{candidate['company']}  |  {candidate['duration']}")
    run_j_comp.italic = True
    run_j_comp.font.color.rgb = RGBColor(80, 80, 80) # Dark Neutral Gray
    p_j_right.paragraph_format.space_after = Pt(0)
    
    # Bullet points for experience details (Pure Black)
    for bullet in candidate["experience_bullets"]:
        p_bullet = doc.add_paragraph(style='List Bullet')
        p_bullet.paragraph_format.space_before = Pt(0)
        p_bullet.paragraph_format.space_after = Pt(3)
        p_bullet.paragraph_format.line_spacing = 1.15
        run_bullet = p_bullet.add_run(bullet)
        run_bullet.font.color.rgb = RGBColor(0, 0, 0)
        
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(8)
    
    # SECTION 4: Education
    h_edu = doc.add_paragraph()
    run_h_edu = h_edu.add_run("EDUCATIONAL QUALIFICATIONS")
    run_h_edu.font.size = Pt(13)
    run_h_edu.bold = True
    run_h_edu.font.color.rgb = RGBColor(0, 0, 0) # Strictly Black
    h_edu.paragraph_format.space_before = Pt(10)
    h_edu.paragraph_format.space_after = Pt(6)
    
    # Education Header Table (Degree on Left, Institution/Year on Right)
    edu_table = doc.add_table(rows=1, cols=2)
    edu_table.autofit = False
    edu_table.columns[0].width = Inches(4.5)
    edu_table.columns[1].width = Inches(2.0)
    
    cell_e_left = edu_table.cell(0, 0)
    cell_e_right = edu_table.cell(0, 1)
    
    p_e_left = cell_e_left.paragraphs[0]
    run_e_deg = p_e_left.add_run(candidate["degree"])
    run_e_deg.bold = True
    run_e_deg.font.size = Pt(12)
    run_e_deg.font.color.rgb = RGBColor(0, 0, 0)
    p_e_left.paragraph_format.space_after = Pt(0)
    
    p_e_right = cell_e_right.paragraphs[0]
    p_e_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_e_inst = p_e_right.add_run(candidate["inst"])
    run_e_inst.italic = True
    run_e_inst.font.color.rgb = RGBColor(80, 80, 80) # Dark Neutral Gray
    p_e_right.paragraph_format.space_after = Pt(0)
    
    # Detailed paragraph under education
    p_edu_desc = doc.add_paragraph()
    p_edu_desc.paragraph_format.space_before = Pt(4)
    p_edu_desc.paragraph_format.space_after = Pt(10)
    p_edu_desc.paragraph_format.line_spacing = 1.15
    run_edu_desc = p_edu_desc.add_run(candidate["edu_details"])
    run_edu_desc.font.italic = True
    run_edu_desc.font.size = Pt(10.5)
    run_edu_desc.font.color.rgb = RGBColor(80, 80, 80) # Dark Neutral Gray
    
    # Save the generated document
    doc.save(filename)
    print(f"Generated beautifully formatted resume: {filename}")

if __name__ == "__main__":
    for c in CANDIDATES:
        generate_docx(c)
    print("\nSuccessfully generated all 10 resumes in the 'cv' folder!")
